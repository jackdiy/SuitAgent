#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown到Word文档转换工具（完整版）
支持自动格式化为法律文书标准格式，现已支持表格转换和格式保留：
- 页面大小：A4 (21cm × 29.7cm)
- 字体：仿宋_GB2312，黑色
- 字号：小四 (12pt)
- 行距：1.5倍
- 首行缩进：2个字符
- 一级标题：小三号，居中加粗，段前段后0.5行
- 其他内容：两端对齐
- 页边距：上下2.54cm，左右3.18cm
- 引号转换：自动将英文引号转换为中文引号
- 页码设置：自动添加页脚页码（格式：1/x，Times New Roman五号）
- 表格转换：支持Markdown表格转换为Word表格，自动设置边框和格式
- 格式支持：支持**加粗**、*斜体*、<u>下划线</u>、~~删除线~~等格式

使用方法：
1. 简单使用：将此脚本放在包含.md文件的文件夹中，运行脚本
2. 指定文件：python md_to_word_converter_complete.py input.md output.docx
3. 使用模板：python md_to_word_converter_complete.py input.md output.docx template.docx
4. 自动模板：程序会自动查找同目录下的.docx文件作为模板（优先使用包含'模板'或'template'的文件）
"""

import os
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import re
import glob
import requests
import base64
import io
from PIL import Image
import tempfile
import subprocess
import shutil
import time
from bs4 import BeautifulSoup

# 全局图片参数（可按需调整）
# 图片显示大小（更大更清晰）：
# - 将默认占比由 70% 提升到 92%
# - 将最大宽度由 10.5cm 提升到 14.2cm（A4 可用宽约 14.64cm）
IMAGE_DISPLAY_RATIO = 0.92         # 相对于页面可用宽度的比例
IMAGE_MAX_DISPLAY_CM = 14.2        # 图片最大显示宽度（cm）
IMAGE_TARGET_DPI = 260             # 目标DPI（用于下采样像素宽度计算）

def get_image_output_path(md_file_path, png_filename):
    """获取图片输出路径，确保目录存在"""
    md_dir = os.path.dirname(os.path.abspath(md_file_path))
    # 基于Markdown文件名创建子目录
    md_filename_base = os.path.splitext(os.path.basename(md_file_path))[0]
    image_dir = os.path.join(md_dir, f"{md_filename_base}_images")
    
    if not os.path.exists(image_dir):
        try:
            os.makedirs(image_dir)
            print(f"📂 创建图片目录: {os.path.relpath(image_dir)}")
        except OSError as e:
            print(f"⚠️ 创建目录失败: {e}")
            return None
            
    return os.path.join(image_dir, png_filename)


def create_word_document(md_file_path, output_path, template_file=None):
    """
    从Markdown文件创建格式化的Word文档
    
    格式要求：
    - 页面大小：A4 (21cm × 29.7cm)
    - 字体：仿宋_GB2312，黑色
    - 字号：小四 (12pt)  
    - 行距：1.5倍
    - 首行缩进：2个字符
    - 一级标题：小三号，居中加粗，段前段后0.5行
    - 其他内容：两端对齐
    - 页边距：上下2.54cm，左右3.18cm
    - 引号转换：自动将英文引号转换为中文引号
    - 页码设置：自动添加页脚页码（格式：1/x）
    - 表格转换：支持Markdown表格转换为Word表格
    - 格式支持：支持各种Markdown格式标记
    """
    
    print(f"📄 正在处理: {md_file_path}")
    
    # 添加引号调试
    debug_quotes_in_file(md_file_path)
    
    # 创建或加载文档
    if template_file and template_file != "none" and os.path.exists(template_file):
        print(f"📋 使用模板文件: {os.path.basename(template_file)}")
        doc = Document(template_file)
        # 清空模板内容
        try:
            # 清空段落和表格
            for paragraph in list(doc.paragraphs):
                if paragraph != doc.paragraphs[0]:
                    p = paragraph._element
                    p.getparent().remove(p)
                else:
                    paragraph.clear()
            
            for table in list(doc.tables):
                t = table._element
                t.getparent().remove(t)
        except Exception as e:
            print(f"⚠️ 清空模板内容失败: {e}")
    else:
        print("📄 创建新文档（不使用模板）")
        doc = Document()
    # 设置默认字体以避免PDF嵌入问题（Normal样式）
    try:
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(10.5)
        normal_style._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        normal_style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        normal_style._element.rPr.rFonts.set(qn('w:cs'), 'Times New Roman')
    except Exception as _:
        pass
    
    # 设置页面大小为A4和页边距
    sections = doc.sections
    for section in sections:
        # 设置页面大小为A4 (21cm × 29.7cm)
        section.page_width = Cm(21.0)      # A4宽度：21cm
        section.page_height = Cm(29.7)     # A4高度：29.7cm
        
        # 设置页边距
        section.top_margin = Cm(2.54)      # 上边距：2.54cm
        section.bottom_margin = Cm(2.54)   # 下边距：2.54cm
        section.left_margin = Cm(3.18)     # 左边距：3.18cm
        section.right_margin = Cm(3.18)    # 右边距：3.18cm
    
    # 读取Markdown文件
    try:
        with open(md_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except UnicodeDecodeError:
        # 如果UTF-8失败，尝试其他编码
        with open(md_file_path, 'r', encoding='gbk') as f:
            content = f.read()
    
    # 按行处理内容，保留原始行结构
    lines = content.split('\n')
    processed_lines = lines  # 保留所有行，包括空行
    
    # 处理表格和图表
    has_body_before_first_h2 = False
    has_seen_h2 = False
    i = 0
    while i < len(processed_lines):
        line = processed_lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # 检查是否是Mermaid图表开始（兼容 ``` mermaid / ```mermaid ）
        if re.match(r'^```\s*mermaid\b', line):
            mermaid_lines = []
            i += 1  # 跳过开始标记
            
            # 收集Mermaid代码
            while i < len(processed_lines) and not processed_lines[i].strip().startswith('```'):
                mermaid_lines.append(processed_lines[i])
                i += 1
            
            if i < len(processed_lines):
                i += 1  # 跳过结束标记
            
            # 处理Mermaid图表
            if mermaid_lines:
                mermaid_code = '\n'.join(mermaid_lines)
                create_mermaid_chart(doc, mermaid_code, md_file_path)
                # 插入了实质内容
                if not has_seen_h2:
                    has_body_before_first_h2 = True
                print(f"✅ 处理Mermaid图表")
            continue
        
        # 处理代码块 ```lang ... ```
        if line.startswith('```'):
            code_lines = []
            language = line[3:].strip()
            i += 1
            while i < len(processed_lines) and not processed_lines[i].strip().startswith('```'):
                code_lines.append(processed_lines[i])
                i += 1
            if i < len(processed_lines):
                i += 1  # 跳过结束标记
            add_code_block(doc, code_lines, language)
            if not has_seen_h2:
                has_body_before_first_h2 = True
            print("✅ 处理代码块")
            continue
        
        # 检查是否是HTML表格开始
        if '<table>' in line.lower():
            html_table_content = []
            table_start = i
            
            # 收集HTML表格的所有内容
            while i < len(processed_lines):
                current_line = processed_lines[i]
                html_table_content.append(current_line)
                if '</table>' in current_line.lower():
                    i += 1
                    break
                i += 1
            
            # 处理HTML表格
            if html_table_content:
                html_content = '\n'.join(html_table_content)
                create_word_table_from_html(doc, html_content)
                if not has_seen_h2:
                    has_body_before_first_h2 = True
            continue
        
        # 检查是否是Markdown表格开始
        if is_table_row(line):
            table_lines = []
            table_start = i
            
            # 收集表格的所有行
            while i < len(processed_lines) and is_table_row(processed_lines[i].strip()):
                table_lines.append(processed_lines[i].strip())
                i += 1
            
            # 处理表格
            if len(table_lines) >= 2:  # 至少要有标题行和分隔行
                create_word_table(doc, table_lines)
                if not has_seen_h2:
                    has_body_before_first_h2 = True
                print(f"✅ 处理Markdown表格: {len(table_lines)} 行")
            continue
        
        # 分割线
        if line in ['---', '***', '___']:
            add_horizontal_line(doc)
            if not has_seen_h2:
                has_body_before_first_h2 = True
            i += 1
            continue
        
        # 任务列表
        if line.startswith('- [ ]') or line.startswith('- [x]') or line.startswith('- [X]'):
            add_task_list(doc, line)
            if not has_seen_h2:
                has_body_before_first_h2 = True
            i += 1
            continue
        
        # 无序列表
        if line.startswith(('- ', '* ', '+ ')):
            add_bullet_list(doc, line)
            if not has_seen_h2:
                has_body_before_first_h2 = True
            i += 1
            continue
        
        # 有序列表
        if re.match(r'^\d+\.\s', line):
            add_numbered_list(doc, line)
            if not has_seen_h2:
                has_body_before_first_h2 = True
            i += 1
            continue
        
        # 引用块（处理多行引用）
        if line.startswith('>'):
            quote_lines = []
            # 收集连续的引用行
            while i < len(lines) and lines[i].startswith('>'):
                quote_content = lines[i][1:].strip()  # 移除 > 符号
                quote_lines.append(quote_content)  # 添加所有行，包括空行
                i += 1
            
            # 将多行引用合并为一个引用块
            if quote_lines:
                full_quote = '\n'.join(quote_lines)
                add_quote(doc, full_quote)
                if not has_seen_h2:
                    has_body_before_first_h2 = True
            continue
            
        # 判断标题级别
        if line.startswith('# '):
            # 一级标题：小三号，居中加粗，段前段后0.5行
            title = line[2:].strip()
            title = convert_quotes_to_chinese(title)  # 转换引号
            p = doc.add_paragraph()
            parse_text_formatting(p, title, title_level=1)  # 使用格式解析处理标题内容
            set_paragraph_format(p, title_level=1)
            
        elif line.startswith('## '):
            # 二级标题：加粗，支持内部格式
            # 在条件满足时，标题前插入一个空行
            if has_seen_h2 or has_body_before_first_h2:
                doc.add_paragraph("")
            title = line[3:].strip()
            title = convert_quotes_to_chinese(title)  # 转换引号
            p = doc.add_paragraph()
            parse_text_formatting(p, title, title_level=2)  # 使用格式解析处理标题内容
            set_paragraph_format(p, title_level=2)
            has_seen_h2 = True
            
        elif line.startswith('### '):
            # 三级标题：不加粗，但支持内部格式
            title = line[4:].strip()
            title = convert_quotes_to_chinese(title)  # 转换引号
            p = doc.add_paragraph()
            parse_text_formatting(p, title, title_level=3)  # 使用格式解析处理标题内容
            set_paragraph_format(p, title_level=3)
            
        elif line.startswith('#### '):
            # 四级标题：不加粗，但支持内部格式
            title = line[5:].strip()
            title = convert_quotes_to_chinese(title)  # 转换引号
            p = doc.add_paragraph()
            parse_text_formatting(p, title, title_level=4)  # 使用格式解析处理标题内容
            set_paragraph_format(p, title_level=4)
            
        else:
            # 正文段落
            if line:
                p = doc.add_paragraph()
                parse_text_formatting(p, line)
                set_paragraph_format(p)
                if not has_seen_h2:
                    has_body_before_first_h2 = True
        
        i += 1
    
    # 添加页码
    add_page_number(doc)
    
    # 保存文档
    doc.save(output_path)
    print(f"✅ Word文档已生成: {output_path}")

def preprocess_mermaid_code(mermaid_code: str) -> str:
    """预处理Mermaid源码，避免Mermaid v11 对标签内Markdown解析导致的
    "Unsupported markdown: list"/"codespan" 等错误。
    - 将行首的 "- ", "* " 项目符号替换为 "• "（兜底）
    - 将编号列表的 "1. " 改为 "1: "（兜底）
    - 将反引号 ` 替换为普通单引号 '，以避免 codespan 报错
    - 重点：对节点标签内部（[...], (...), ((...)), {...}, >...], ["..."], ("...")) 的起始列表标记进行替换
    该处理为无害替换，不影响边、样式等语句。
    """
    import re

    s = mermaid_code

    # 反引号替换，避免 codespan 被解析
    s = s.replace("`", "'")

    # 1) 针对节点标签内部：有序列表 1. -> 1:
    def _repl_number_dot(m: re.Match) -> str:
        brace = m.group('brace')
        quote = m.group('quote') or ''
        num = m.group('num')
        return f"{brace}{quote}{num}: "

    s = re.sub(r"(?m)(?P<brace>[\[\({\>])(?P<quote>\"?\s*)(?P<num>\d+)\.\s", _repl_number_dot, s)

    # 2) 针对节点标签内部：无序列表 - / * -> •
    def _repl_bullet(m: re.Match) -> str:
        brace = m.group('brace')
        quote = m.group('quote') or ''
        return f"{brace}{quote}• "

    s = re.sub(r"(?m)(?P<brace>[\[\({\>])(?P<quote>\"?\s*)[-*]\s", _repl_bullet, s)

    # 3) 兜底：整行以列表开头的情况（极少出现在Mermaid内，但保留以防万一）
    s = re.sub(r"(?m)^(\s*)-\s+", r"\1• ", s)
    s = re.sub(r"(?m)^(\s*)\*\s+", r"\1• ", s)
    s = re.sub(r"(?m)^(\s*)(\d+)\.\s+", r"\1\2: ", s)

    return s

def create_mermaid_chart(doc, mermaid_code, md_file_path):
    """将Mermaid图表转换为图片并插入Word文档（本地渲染优先）"""

    # 预处理，规避 Mermaid 11 对列表/反引号的 Markdown 解析造成的报错
    mermaid_code = preprocess_mermaid_code(mermaid_code)

    # 首先尝试本地渲染
    local_success = try_local_mermaid_render(doc, mermaid_code, md_file_path)
    if local_success:
        return

    # 仅使用本地渲染：失败则改为文本，不再尝试在线服务
    print("⚠️ 本地渲染失败，已禁用在线服务，使用文本替代")
    create_fallback_text(doc, mermaid_code)

def try_local_mermaid_render(doc, mermaid_code, md_file_path):
    """尝试使用本地mermaid-cli渲染图表"""
    
    # 为Mermaid文件和输出图片准备路径
    timestamp = str(int(time.time() * 1000))
    mmd_filename = f"mermaid-src-{timestamp}.mmd"
    png_filename = f"mermaid-chart-{timestamp}.png"
    
    # 获取保存图片的最终路径
    output_png_path = get_image_output_path(md_file_path, png_filename)
    if not output_png_path:
        print("⚠️ 无法获取图片输出路径，跳过本地渲染。")
        return False
        
    # 临时文件放在脚本所在目录，避免 cwd 不一致导致路径问题
    script_dir = os.path.dirname(os.path.abspath(__file__))
    temp_mmd_path = os.path.join(script_dir, mmd_filename)

    try:
        print("🖥️ 尝试本地Mermaid渲染...")
        
        # 创建临时的.mmd文件
        with open(temp_mmd_path, 'w', encoding='utf-8') as f:
            f.write(mermaid_code)
        
        # 检查 mmdc 命令：优先环境变量 MMDCCMD，其次脚本同目录 node_modules，再其次系统 PATH
        mmdc_env = os.environ.get('MMDCCMD', '').strip()
        mmdc_path = mmdc_env if mmdc_env else os.path.join(script_dir, "node_modules", ".bin", "mmdc")
        if not os.path.exists(mmdc_path):
            mmdc_path = shutil.which("mmdc") or ""
        if not mmdc_path:
            print("⚠️ 本地 mmdc 命令未找到（已跳过本地渲染）")
            return False
        
        # 使用mmdc命令生成高分辨率PNG图片
        # 绝对路径，配置文件若存在则使用
        abs_in = os.path.abspath(temp_mmd_path)
        abs_out = os.path.abspath(output_png_path)
        cfg = os.path.join(script_dir, "mermaid-config.json")
        cmd = [mmdc_path, "-i", abs_in, "-o", abs_out, "-t", "neutral", "-w", "2200", "-H", "1500", "--scale", "2.0"]
        if os.path.exists(cfg):
            cmd.extend(["-c", cfg])
        
        print(f"🔧 执行命令: {' '.join(cmd)}")
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        
        if result.returncode != 0:
            print(f"⚠️ mmdc 命令执行失败: {result.stderr}")
            return False
        
        # 检查生成的PNG文件是否存在
        if not os.path.exists(output_png_path):
            print("⚠️ PNG文件未生成")
            return False
        
        # 加载图片并插入Word
        image = Image.open(output_png_path)
        insert_image_to_word(doc, image)
        
        print(f"✅ 本地Mermaid图表渲染成功！图片已保存至: {os.path.relpath(output_png_path)}")
        return True
        
    except subprocess.TimeoutExpired:
        print("⚠️ mmdc命令执行超时")
        return False
    except Exception as e:
        print(f"⚠️ 本地渲染失败: {e}")
        return False
    finally:
        # 无论成功与否，都清理临时的mmd文件
        if os.path.exists(temp_mmd_path):
            try:
                os.unlink(temp_mmd_path)
            except:
                pass

def _postprocess_image_for_word(image, target_display_cm, target_dpi=IMAGE_TARGET_DPI):
    """根据目标显示宽度与DPI对图像进行高质量下采样，控制体积并保持清晰度"""
    try:
        # 目标像素宽度 = 目标显示英寸 * 目标DPI
        target_inches = float(target_display_cm) / 2.54
        target_px_width = max(1, int(target_inches * target_dpi))
        if image.width > target_px_width:
            new_height = int(image.height * (target_px_width / image.width))
            image = image.resize((target_px_width, new_height), Image.LANCZOS)
    except Exception:
        pass
    return image

def insert_image_to_word(doc, image):
    """将PIL图片对象插入Word文档"""
    
    # 保存到临时文件
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
        # 计算页面可用宽度与目标插入宽度（限制较小以保持版面协调）
        available_width_cm = 21.0 - 3.18 * 2
        target_display_cm = min(available_width_cm * IMAGE_DISPLAY_RATIO, IMAGE_MAX_DISPLAY_CM)
        # 按目标DPI对图像进行高质量下采样（清晰但不臃肿）
        image = _postprocess_image_for_word(image, target_display_cm, target_dpi=IMAGE_TARGET_DPI)
        # 使用高压缩PNG保存，进一步降低体积
        try:
            image.save(temp_file.name, format='PNG', optimize=True, compress_level=9)
        except Exception:
            image.save(temp_file.name, format='PNG', optimize=True)
        temp_filename = temp_file.name
    
    try:
        # 在Word中插入图片
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 统一的目标显示宽度，保持版面一致性
        img_width_cm = target_display_cm
        
        run = paragraph.add_run()
        run.add_picture(temp_filename, width=Cm(img_width_cm))
        
    finally:
        # 删除临时文件
        try:
            os.unlink(temp_filename)
        except:
            pass

def create_fallback_text(doc, mermaid_code):
    """创建后备文本方案"""
    # 解析图表类型并创建简化版本
    if 'graph' in mermaid_code.lower():
        create_simple_diagram_text(doc, mermaid_code)
    elif 'pie' in mermaid_code.lower():
        create_simple_pie_text(doc, mermaid_code)
    elif 'gantt' in mermaid_code.lower():
        create_simple_gantt_text(doc, mermaid_code)
    else:
        # 默认处理
        p = doc.add_paragraph()
        run = p.add_run("【图表内容】")
        run.bold = True
        p.add_run("\n" + mermaid_code)
        set_paragraph_format(p)

def create_simple_diagram_text(doc, mermaid_code):
    """创建简化的图表文本描述"""
    p = doc.add_paragraph()
    run = p.add_run("【流程图】")
    run.bold = True
    
    # 解析节点和连接关系
    lines = mermaid_code.split('\n')
    nodes = {}
    connections = []
    
    for line in lines:
        line = line.strip()
        if '-->' in line or '->' in line:
            parts = line.split('-->' if '-->' in line else '->')
            if len(parts) == 2:
                from_node = parts[0].strip()
                to_node = parts[1].strip()
                connections.append(f"{from_node} → {to_node}")
        elif '[' in line and ']' in line:
            # 解析节点定义
            match = re.search(r'(\w+)\["([^"]+)"\]', line)
            if match:
                node_id, node_text = match.groups()
                nodes[node_id] = node_text
    
    # 添加解析结果
    if connections:
        p.add_run("\n主要流程:")
        for conn in connections[:8]:  # 最多显示8个连接
            p.add_run(f"\n• {conn}")
    
    set_paragraph_format(p)

def create_simple_pie_text(doc, mermaid_code):
    """创建简化的饼图文本描述"""
    p = doc.add_paragraph()
    run = p.add_run("【数据分析】")
    run.bold = True
    
    # 解析饼图数据
    lines = mermaid_code.split('\n')
    for line in lines:
        if ':' in line and '"' in line:
            # 解析数据项
            match = re.search(r'"([^"]+)"\s*:\s*(\d+(?:\.\d+)?)', line)
            if match:
                label, value = match.groups()
                p.add_run(f"\n• {label}: {value}")
    
    set_paragraph_format(p)

def create_simple_gantt_text(doc, mermaid_code):
    """创建简化的甘特图文本描述"""
    p = doc.add_paragraph()
    run = p.add_run("【时间安排】")
    run.bold = True
    
    # 解析甘特图任务
    lines = mermaid_code.split('\n')
    current_section = ""
    
    for line in lines:
        line = line.strip()
        if line.startswith('section '):
            current_section = line.replace('section ', '')
            p.add_run(f"\n\n{current_section}:")
        elif ':' in line and not line.startswith('title'):
            # 解析任务
            task = line.split(':')[0].strip()
            p.add_run(f"\n• {task}")
    
    set_paragraph_format(p)

def is_separator_line(line):
    """判断是否是表格分隔行。分隔行必须包含'-'，且只能包含'|', '-', ':', ' '等符号。"""
    line = line.strip()
    if not line or '-' not in line:
        return False
    return all(c in '|-: 	' for c in line)


def is_table_row(line):
    """判断是否是表格行"""
    if not line or not line.strip():
        return False
    
    line = line.strip()
    
    # 检查是否是分隔行
    if is_separator_line(line):
        return True
    
    # 检查是否是数据行（包含 |）
    # 这里的逻辑保持宽松，依赖于主循环中对其他块级元素的优先判断
    if '|' in line:
        return True
    
    return False

def create_word_table(doc, table_lines):
    """从Markdown表格行创建Word表格"""
    
    if len(table_lines) < 2:
        return
    
    # 解析表格数据
    rows_data = []
    header_row = None
    
    for i, line in enumerate(table_lines):
        # 跳过分隔行（包含横线的行）
        if is_separator_line(line):
            continue
        
        # 解析单元格
        cells = parse_table_row(line)
        if cells:
            if header_row is None:
                header_row = cells
            else:
                rows_data.append(cells)
    
    if not header_row:
        return
    
    # 确定列数
    max_cols = len(header_row)
    for row in rows_data:
        max_cols = max(max_cols, len(row))
    
    # 创建Word表格
    total_rows = 1 + len(rows_data)  # 标题行 + 数据行
    table = doc.add_table(rows=total_rows, cols=max_cols)
    
    # 设置表格样式
    # table.style = 'Table Grid'  # 注释掉可能不存在的样式
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 统一设置边框和内边距、行高等
    try:
        tbl = table._tbl
        borders_xml = '''
        <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        </w:tblBorders>
        '''
        tbl.tblPr.append(parse_xml(borders_xml))
        cell_margins_xml = '''
        <w:tblCellMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:w="30" w:type="dxa"/>
            <w:left w:w="60" w:type="dxa"/>
            <w:bottom w:w="30" w:type="dxa"/>
            <w:right w:w="60" w:type="dxa"/>
        </w:tblCellMar>
        '''
        tbl.tblPr.append(parse_xml(cell_margins_xml))
    except Exception:
        pass
    # 行高与段落行距统一
    try:
        for row in table.rows:
            row.height = Cm(0.8)
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    pf = paragraph.paragraph_format
                    pf.line_spacing = 1.2
                    pf.space_before = Pt(2)
                    pf.space_after = Pt(2)
    except Exception:
        pass
    
    # 填充标题行
    header_cells = table.rows[0].cells
    for j, cell_text in enumerate(header_row):
        if j < len(header_cells):
            cell = header_cells[j]
            # 处理表格单元格中的格式
            if contains_markdown_formatting(cell_text.strip()):
                parse_table_cell_formatting(cell, cell_text.strip(), is_header=True)
            else:
                cell.text = convert_quotes_to_chinese(cell_text.strip())
                set_table_cell_format(cell, is_header=True)
    
    # 填充数据行
    for i, row_data in enumerate(rows_data):
        if i + 1 < len(table.rows):
            row_cells = table.rows[i + 1].cells
            for j, cell_text in enumerate(row_data):
                if j < len(row_cells):
                    cell = row_cells[j]
                    # 处理表格单元格中的格式
                    if contains_markdown_formatting(cell_text.strip()):
                        parse_table_cell_formatting(cell, cell_text.strip(), is_header=False)
                    else:
                        cell.text = convert_quotes_to_chinese(cell_text.strip())
                        set_table_cell_format(cell, is_header=False)
    
    # 调整列宽
    adjust_table_column_width(table)

def parse_table_row(line):
    """解析表格行，提取单元格内容"""
    if not line or not line.strip():
        return []
    
    line = line.strip()
    
    # 移除开头和结尾的 |
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    
    # 分割单元格
    cells = [cell.strip() for cell in line.split('|')]
    
    # 过滤掉空单元格（但保留有意义的空单元格）
    return cells

def contains_markdown_formatting(text):
    """检查文本是否包含Markdown格式标记"""
    format_patterns = [
        r'\*\*\*.*?\*\*\*',  # 加粗斜体
        r'\*\*.*?\*\*',      # 加粗
        r'\*.*?\*',          # 斜体
        r'___.*?___',        # 加粗斜体
        r'__.*?__',          # 加粗
        r'_.*?_',            # 斜体
        r'<u>.*?</u>',       # 下划线
        r'~~.*?~~',          # 删除线
        r'`.*?`',            # 行内代码
        r'<br\s*/?>',       # 换行标签
        r'\$.*?\$',         # LaTeX数学公式
    ]
    
    for pattern in format_patterns:
        if re.search(pattern, text):
            return True
    return False

def parse_table_cell_formatting(cell, text, is_header=False):
    """解析表格单元格中的格式化文本"""
    # 清空单元格
    cell.text = ""
    
    # 转换引号
    text = convert_quotes_to_chinese(text)
    
    # 支持<br>换行：拆分后逐段处理
    parts_by_br = re.split(r'<br\s*/?>', text, flags=re.IGNORECASE)
    
    # 解析格式
    format_patterns = [
        (r'\*\*\*(.*?)\*\*\*', {'bold': True, 'italic': True}),
        (r'___(.*?)___', {'bold': True, 'italic': True}),
        (r'\*\*(.*?)\*\*', {'bold': True}),
        (r'__(.*?)__', {'bold': True}),
        (r'(?<!\*)\*([^*\n]+?)\*(?!\*)', {'italic': True}),
        (r'(?<!_)_([^_\n]+?)_(?!_)', {'italic': True}),
        (r'<u>(.*?)</u>', {'underline': True}),
        (r'~~(.*?)~~', {'strikethrough': True}),
        (r'`([^`\n]+)`', {'code': True}),
        (r'\$([^$\n]+?)\$', {'math': True}),  # LaTeX数学公式支持
    ]
    
    for idx, segment in enumerate(parts_by_br):
        if idx > 0:
            cell.paragraphs[0].add_run().add_break()
        text_parts = parse_formatted_text(segment, format_patterns)
        for part_text, formats in text_parts:
            if part_text:  # 只有非空文本才创建run
                run = cell.paragraphs[0].add_run(part_text)
                set_table_run_format(run, formats, is_header)

def set_table_run_format(run, formats, is_header=False):
    """设置表格单元格run格式"""
    font = run.font
    font.name = 'Times New Roman'  # 默认英文字体
    font.size = Pt(10.5)  # 表格使用五号字体
    font.color.rgb = RGBColor(0, 0, 0)
    font.bold = is_header  # 标题行基础加粗
    
    # 设置字体映射：英文和数字用Times New Roman，中文用仿宋_GB2312
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')     # ASCII字符（英文字母、数字、标点）
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')     # 高位ANSI字符
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')       # 东亚字符（中文）
    run._element.rPr.rFonts.set(qn('w:cs'), 'Times New Roman')        # 复杂脚本字符
    
    # 应用Markdown格式
    if formats.get('bold', False):
        font.bold = True
    if formats.get('italic', False):
        font.italic = True
    if formats.get('underline', False):
        font.underline = True
    if formats.get('strikethrough', False):
        font.strike = True
    if formats.get('code', False):
        # 表格中代码使用Times New Roman，稍小字号
        font.name = 'Times New Roman'
        font.size = Pt(9)
        font.color.rgb = RGBColor(51, 51, 51)
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        return
    if formats.get('math', False):
        # 表格中数学公式使用Times New Roman，斜体，深蓝色
        font.name = 'Times New Roman'
        font.size = Pt(10)
        font.italic = True
        font.color.rgb = RGBColor(0, 0, 139)  # 深蓝色
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        return

def set_table_cell_format(cell, is_header=False):
    """设置表格单元格格式"""
    
    # 设置段落格式
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0  # 表格内使用单倍行距
        
        # 设置文字格式
        for run in paragraph.runs:
            font = run.font
            font.name = '仿宋_GB2312'
            font.size = Pt(10.5)  # 表格使用五号字体
            font.color.rgb = RGBColor(0, 0, 0)
            font.bold = is_header  # 标题行加粗
            
            # 设置中文字体
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

def adjust_table_column_width(table):
    """调整表格列宽"""
    try:
        # 获取表格总宽度（页面宽度减去页边距）
        available_width = Cm(21.0 - 3.18 * 2)  # A4宽度减去左右页边距
        
        # 平均分配列宽
        col_count = len(table.columns)
        if col_count > 0:
            col_width = int(available_width / col_count)  # 转换为整数
            for column in table.columns:
                column.width = col_width
    except Exception as e:
        print(f"⚠️  表格列宽调整失败: {e}")

def parse_html_table(html_content):
    """解析HTML表格内容，返回表格数据"""
    try:
        soup = BeautifulSoup(html_content, 'html.parser')
        table = soup.find('table')
        if not table:
            return None
        
        rows_data = []
        for tr in table.find_all('tr'):
            row_cells = []
            for cell in tr.find_all(['td', 'th']):
                # 获取单元格文本内容，保留基本格式
                cell_text = cell.get_text(strip=True)
                row_cells.append(cell_text)
            if row_cells:  # 只添加非空行
                rows_data.append(row_cells)
        
        return rows_data
    except Exception as e:
        print(f"⚠️  HTML表格解析失败: {e}")
        return None

def create_word_table_from_html(doc, html_content):
    """从HTML表格创建Word表格"""
    rows_data = parse_html_table(html_content)
    if not rows_data or len(rows_data) < 1:
        print("⚠️  HTML表格数据为空或格式不正确")
        return
    
    # 创建Word表格
    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    
    # 设置表格样式（使用自定义边框而不是Table Grid样式）
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 设置表格边框和单元格边距
    try:
        tbl = table._tbl
        borders_xml = '''
        <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        </w:tblBorders>
        '''
        tbl.tblPr.append(parse_xml(borders_xml))
        cell_margins_xml = '''
        <w:tblCellMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:w="30" w:type="dxa"/>
            <w:left w:w="60" w:type="dxa"/>
            <w:bottom w:w="30" w:type="dxa"/>
            <w:right w:w="60" w:type="dxa"/>
        </w:tblCellMar>
        '''
        tbl.tblPr.append(parse_xml(cell_margins_xml))
    except Exception:
        pass
    
    # 设置行高和单元格对齐
    try:
        for row in table.rows:
            row.height = Cm(0.8)
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    pf = paragraph.paragraph_format
                    pf.line_spacing = 1.2
                    pf.space_before = Pt(2)
                    pf.space_after = Pt(2)
    except Exception:
        pass
    
    # 填充表格数据
    for i, row_data in enumerate(rows_data):
        if i < len(table.rows):
            row_cells = table.rows[i].cells
            for j, cell_text in enumerate(row_data):
                if j < len(row_cells):
                    cell = row_cells[j]
                    cell.text = convert_quotes_to_chinese(cell_text.strip())
                    # 第一行作为标题行处理
                    set_table_cell_format(cell, is_header=(i == 0))
    
    # 调整列宽
    adjust_table_column_width(table)
    print(f"✅ 处理HTML表格: {len(rows_data)} 行")

def convert_quotes_to_chinese(text):
    """将英文引号转换为中文引号（改进版）
    规则：
    - 将直双引号 " 转为中文开/闭引号 “ ”（使用简单语境+切换判断）
    - 将直单引号 ' 转为中文开/闭引号 ‘ ’，但保留英文缩写/所有格中的撇号（如 don't, John's）
    - 避免转换代码片段中的引号（由反引号 ` 包裹）
    """
    if not text:
        return text

    original_text = text

    # 若无需要处理的引号，直接返回
    if ('"' not in text) and ("'" not in text):
        return text

    result = []
    i = 0
    in_code = False  # 是否处于 `code` 片段中

    while i < len(text):
        ch = text[i]

        # 处理反引号包裹的代码片段，保持原样
        if ch == '`':
            # 统计连续反引号的数量（支持 ``` 块 及 ` 行内`）
            j = i + 1
            while j < len(text) and text[j] == '`':
                j += 1
            backtick_count = j - i
            result.append('`' * backtick_count)
            in_code = not in_code  # 简化处理：遇到成组反引号时翻转状态
            i = j
            continue

        if in_code:
            # 代码片段内不做引号替换
            result.append(ch)
            i += 1
            continue

        if ch == '"':
            # 判断前一个非空白字符，以推断开/闭引号
            k = len(result) - 1
            prev_char = None
            while k >= 0:
                pc = result[k]
                if not pc.isspace():
                    prev_char = pc
                    break
                k -= 1
            # 若前面为空/是开括号/标点，则更可能是开引号
            if prev_char is None or prev_char in '([{<（【《“‘\t\n "\'\-—:;,.!?、，。；：！？”）〉》»…':
                result.append('“')
            else:
                result.append('”')
            i += 1
            continue

        if ch == "'":
            # 保留英文缩写/所有格中的撇号：字母-撇号-字母
            prev_c = text[i - 1] if i > 0 else ''
            next_c = text[i + 1] if i + 1 < len(text) else ''
            if prev_c.isalpha() and next_c.isalpha():
                result.append("'")
                i += 1
                continue

            # 判断前一个非空白字符，以推断开/闭单引号
            k = len(result) - 1
            prev_char = None
            while k >= 0:
                pc = result[k]
                if not pc.isspace():
                    prev_char = pc
                    break
                k -= 1
            if prev_char is None or prev_char in '([{<（【《“\t\n "\'\-—:;,.!?、，。；：！？”）〉》»…':
                result.append('‘')
            else:
                result.append('’')
            i += 1
            continue

        # 其它字符保持
        result.append(ch)
        i += 1

    text = ''.join(result)

    if text != original_text:
        print(f"✅ 引号转换: {original_text} → {text}")

    return text

def add_page_number(doc):
    """添加页码，格式为 '页码/总页数'"""
    
    try:
        # 获取文档的第一个节
        section = doc.sections[0]
        
        # 获取页脚
        footer = section.footer
        
        # 清空现有页脚内容
        for para in footer.paragraphs:
            para.clear()
        
        # 如果没有段落，添加一个
        if not footer.paragraphs:
            footer_para = footer.add_paragraph()
        else:
            footer_para = footer.paragraphs[0]
        
        # 设置段落居中对齐
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 创建页码文本的XML
        from docx.oxml.shared import qn
        
        # 添加当前页码字段
        run = footer_para.add_run()
        
        # 创建页码字段
        fld_char_begin = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        run._r.append(fld_char_begin)
        
        instr_text = parse_xml(r'<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE </w:instrText>')
        run._r.append(instr_text)
        
        fld_char_end = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        run._r.append(fld_char_end)
        
        # 添加分隔符
        sep_run = footer_para.add_run("/")
        
        # 添加总页数字段
        total_run = footer_para.add_run()
        
        fld_char_begin2 = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        total_run._r.append(fld_char_begin2)
        
        instr_text2 = parse_xml(r'<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> NUMPAGES </w:instrText>')
        total_run._r.append(instr_text2)
        
        fld_char_end2 = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        total_run._r.append(fld_char_end2)
        
        # 设置所有run的字体格式为Times New Roman五号
        for run in footer_para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)  # 五号字体
            run.font.color.rgb = RGBColor(0, 0, 0)
            # 设置西文字体为Times New Roman
            run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
            run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            
    except Exception as e:
        print(f"⚠️  页码添加失败，将跳过页码设置: {e}")
        pass

def parse_text_formatting(paragraph, text, title_level=0, is_quote=False):
    """解析文本格式（支持加粗、斜体、下划线，转换引号为中文）"""
    
    # 转换英文引号为中文引号
    text = convert_quotes_to_chinese(text)
    
    # 先处理<br>标签为段内换行
    segments = re.split(r'<br\s*/?>', text, flags=re.IGNORECASE)
    
    # 使用正则表达式解析所有格式标记
    format_patterns = [
        (r'\*\*\*(.*?)\*\*\*', {'bold': True, 'italic': True}),
        (r'___(.*?)___', {'bold': True, 'italic': True}),
        (r'\*\*(.*?)\*\*', {'bold': True}),
        (r'__(.*?)__', {'bold': True}),
        (r'(?<!\*)\*([^*\n]+?)\*(?!\*)', {'italic': True}),
        (r'(?<!_)_([^_\n]+?)_(?!_)', {'italic': True}),
        (r'<u>(.*?)</u>', {'underline': True}),
        (r'~~(.*?)~~', {'strikethrough': True}),
        (r'`([^`\n]+)`', {'code': True}),
        (r'\$([^$\n]+?)\$', {'math': True}),  # LaTeX数学公式支持
    ]
    
    for idx, segment in enumerate(segments):
        text_parts = parse_formatted_text(segment, format_patterns)
        for part_text, formats in text_parts:
            if part_text:  # 只有非空文本才创建run
                run = paragraph.add_run(part_text)
                set_run_format_with_styles(run, formats, title_level=title_level, is_quote=is_quote)
        if idx < len(segments) - 1:
            paragraph.add_run().add_break()

def parse_formatted_text(text, format_patterns):
    """解析带格式的文本，返回(文本, 格式)的列表"""
    
    if not text:
        return []
    
    parts = []
    current_pos = 0
    
    # 查找所有格式标记的位置
    all_matches = []
    for pattern, format_dict in format_patterns:
        for match in re.finditer(pattern, text):
            all_matches.append({
                'start': match.start(),
                'end': match.end(),
                'text': match.group(1),
                'format': format_dict,
                'full_match': match.group(0)
            })
    
    # 按开始位置排序
    all_matches.sort(key=lambda x: x['start'])
    
    # 处理重叠的匹配（选择最长的匹配）
    filtered_matches = []
    for match in all_matches:
        # 检查是否与已有匹配重叠
        overlap = False
        for existing in filtered_matches:
            if (match['start'] < existing['end'] and match['end'] > existing['start']):
                # 有重叠，选择更长的匹配
                if len(match['full_match']) > len(existing['full_match']):
                    filtered_matches.remove(existing)
                    filtered_matches.append(match)
                overlap = True
                break
        if not overlap:
            filtered_matches.append(match)
    
    # 重新按位置排序
    filtered_matches.sort(key=lambda x: x['start'])
    
    # 构建文本部分列表
    for match in filtered_matches:
        # 添加前面的普通文本
        if current_pos < match['start']:
            normal_text = text[current_pos:match['start']]
            if normal_text:
                parts.append((normal_text, {}))
        
        # 添加格式化文本
        parts.append((match['text'], match['format']))
        current_pos = match['end']
    
    # 添加剩余的普通文本
    if current_pos < len(text):
        remaining_text = text[current_pos:]
        if remaining_text:
            parts.append((remaining_text, {}))
    
    # 如果没有找到任何格式，返回整个文本作为普通文本
    if not parts:
        parts.append((text, {}))
    
    return parts

def set_run_format(run, title_level=0):
    """设置文本运行格式（基础版本，用于标题）"""
    font = run.font
    font.name = 'Times New Roman'  # 默认英文字体
    font.color.rgb = RGBColor(0, 0, 0)  # 黑色
    font.bold = False  # 默认不加粗
    font.italic = False  # 不斜体
    font.underline = False  # 不下划线
    
    # 设置字体映射：英文和数字用Times New Roman，中文用仿宋_GB2312
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')     # ASCII字符（英文字母、数字、标点）
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')     # 高位ANSI字符
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')       # 东亚字符（中文）
    run._element.rPr.rFonts.set(qn('w:cs'), 'Times New Roman')        # 复杂脚本字符
    
    if title_level == 1:
        # 一级标题：小三号，加粗
        font.size = Pt(15)  # 小三号字体
        font.bold = True
    elif title_level == 2:
        # 二级标题：小四号，加粗
        font.size = Pt(12)  # 小四号字体
        font.bold = True
    else:
        # 其他：小四号，不加粗
        font.size = Pt(12)  # 小四号字体
        font.bold = False

def set_run_format_with_styles(run, formats, title_level=0, is_quote=False):
    """设置文本运行格式（支持多种样式）"""
    font = run.font
    font.name = 'Times New Roman'  # 默认英文字体
    font.color.rgb = RGBColor(0, 0, 0)  # 黑色
    
    # 设置字体映射：英文和数字用Times New Roman，中文用仿宋_GB2312
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')     # ASCII字符（英文字母、数字、标点）
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')     # 高位ANSI字符
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')       # 东亚字符（中文）
    run._element.rPr.rFonts.set(qn('w:cs'), 'Times New Roman')        # 复杂脚本字符
    
    # 设置基础格式
    if title_level == 1:
        # 一级标题：小三号，加粗
        font.size = Pt(15)  # 小三号字体
        font.bold = True
    elif title_level == 2:
        # 二级标题：小四号，加粗
        font.size = Pt(12)  # 小四号字体
        font.bold = True
    elif is_quote:
        # 引用：小五号
        font.size = Pt(9)  # 小五号字体
        font.bold = False
    else:
        # 其他：小四号
        font.size = Pt(12)  # 小四号字体
        font.bold = False
    
    # 应用Markdown格式（这里是关键）
    if formats.get('code', False):
        # 代码字体使用 Times New Roman，避免PDF嵌入问题
        font.name = 'Times New Roman'
        font.size = Pt(10)
        font.color.rgb = RGBColor(51, 51, 51)
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    elif formats.get('math', False):
        # LaTeX数学公式：使用Times New Roman，稍小字号，斜体
        font.name = 'Times New Roman'
        font.size = Pt(11)  # 比正文稍小
        font.italic = True
        font.color.rgb = RGBColor(0, 0, 139)  # 深蓝色
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    else:
        if formats.get('bold', False):
            font.bold = True
        if formats.get('italic', False):
            font.italic = True
        if formats.get('underline', False):
            font.underline = True
        if formats.get('strikethrough', False):
            font.strike = True

def add_horizontal_line(doc):
    """添加分割线"""
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('─' * 55)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 128, 128)

def add_task_list(doc, line):
    """添加任务列表"""
    is_checked = line.startswith(('- [x]', '- [X]'))
    text = line[5:].strip()
    p = doc.add_paragraph()
    checkbox_run = p.add_run('☑ ' if is_checked else '☐ ')
    set_run_format_with_styles(checkbox_run, {}, title_level=0)
    parse_text_formatting(p, text)
    set_paragraph_format(p)

def add_bullet_list(doc, line):
    """添加无序列表"""
    text = line[2:].strip()
    p = doc.add_paragraph()
    bullet_run = p.add_run('• ')
    set_run_format_with_styles(bullet_run, {}, title_level=0)
    parse_text_formatting(p, text)
    set_paragraph_format(p)


def add_numbered_list(doc, line):
    """添加有序列表（保持原样输出）"""
    p = doc.add_paragraph()
    parse_text_formatting(p, line)
    set_paragraph_format(p)

def add_quote(doc, text):
    """添加带有着重底色的引用，并处理内部列表和多行文本"""
    # 按换行符分割文本，处理每一行
    lines = text.split('\n')
    
    for line_index, line in enumerate(lines):
        if not line.strip():  # 处理空行
            # 添加空段落来保持间距
            p = doc.add_paragraph()
            set_paragraph_format(p, is_quote=True)
            continue
            
        p = doc.add_paragraph()
        
        # 设置段落底色
        from docx.oxml.shared import OxmlElement
        from docx.oxml.ns import qn
        
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'EAEAEA') # 稍深的灰色
        pPr.append(shd)
        
        # 设置左侧缩进
        p.paragraph_format.left_indent = Inches(0.2)
        
        # 检查并处理列表标记
        bullet_match = re.match(r'^\s*([-*+])\s+', line)
        number_match = re.match(r'^\s*(\d+\.)\s+', line)
        
        list_marker_run = None
        
        if bullet_match:
            # 无序列表，使用符号并添加缩进
            indent_and_bullet = '    •  ' # 4个空格缩进 + bullet
            list_marker_run = p.add_run(indent_and_bullet)
            line = line[bullet_match.end():]
        elif number_match:
            # 有序列表，使用数字并添加缩进
            indent_and_number = f'    {number_match.group(1)} '
            list_marker_run = p.add_run(indent_and_number)
            line = line[number_match.end():]

        # 为列表标记设置统一格式
        if list_marker_run:
            set_run_format_with_styles(list_marker_run, {}, is_quote=True)

        # 添加并解析文本内容
        parse_text_formatting(p, line, is_quote=True)
        
        # 调整段落格式
        set_paragraph_format(p, is_quote=True)

def add_code_block(doc, code_lines, language):
    """添加代码块（Times New Roman，避免PDF嵌入问题）"""
    if language:
        lang_p = doc.add_paragraph()
        lang_run = lang_p.add_run(f"[{language}]")
        lang_run.font.name = 'Times New Roman'
        lang_run.font.size = Pt(10)
        lang_run.font.color.rgb = RGBColor(128, 128, 128)
    for code_line in code_lines:
        p = doc.add_paragraph()
        run = p.add_run(code_line or ' ')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(51, 51, 51)
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.line_spacing = 1.2

def set_paragraph_format(paragraph, title_level=0, is_quote=False):
    """设置段落格式"""
    
    # 设置段落格式
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1.5  # 1.5倍行距
    
    if title_level == 1:
        # 一级标题：居中，段前段后0.5行，不缩进
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph_format.space_before = Pt(6)  # 0.5行约6pt
        paragraph_format.space_after = Pt(6)   # 0.5行约6pt
        paragraph_format.first_line_indent = Pt(0)  # 一级标题不缩进
    elif is_quote:
        # 引用：两端对齐，无首行缩进
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.first_line_indent = Pt(0)
    else:
        # 其他段落（二级标题、三级标题、正文）：首行缩进2个字符，两端对齐
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.first_line_indent = Pt(24)  # 2个字符缩进（12pt * 2）
    
    # 确保所有runs都有正确的格式
    for run in paragraph.runs:
        if not hasattr(run.font, 'name') or not run.font.name:
            set_run_format(run, title_level)

def find_template_file():
    """查找模板文件"""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    docx_files = glob.glob(os.path.join(script_dir, "*.docx"))
    
    for docx_file in docx_files:
        filename = os.path.basename(docx_file).lower()
        if not any(keyword in filename for keyword in ['完整版', 'test', 'output', '输出']):
            if '模板' in filename or 'template' in filename:
                return docx_file
    
    return docx_files[0] if docx_files else None


def find_md_files():
    """查找脚本所在目录下的所有 .md 文件（与运行目录无关）"""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_files = glob.glob(os.path.join(script_dir, "*.md"))
    return md_files

def generate_output_filename(md_file):
    """根据输入文件名生成输出文件名"""
    base_name = os.path.splitext(md_file)[0]
    return f"{base_name}_完整版.docx"

def main():
    """主函数"""
    
    print("=" * 60)
    print("📝 Markdown到Word转换工具（完整版）")
    print("支持表格转换和格式保留：仿宋_GB2312, 黑色, 小四, 1.5倍行距")
    print("=" * 60)
    
    # 查找模板文件
    template_file = find_template_file()
    
    if template_file:
        print(f"🎨 发现模板文件: {os.path.basename(template_file)}")
    else:
        print("📄 未发现模板文件，将创建新文档")
    
    # 检查命令行参数
    script_dir = os.path.dirname(os.path.abspath(__file__))
    if len(sys.argv) >= 3:
        # 用户指定了输入和输出文件
        md_file = sys.argv[1]
        # 若传入的是相对路径且当前工作目录不同，尝试以脚本目录补全
        if not os.path.isabs(md_file):
            alt = os.path.join(script_dir, md_file)
            if os.path.exists(alt):
                md_file = alt
        output_file = sys.argv[2]
        template = sys.argv[3] if len(sys.argv) > 3 else template_file
        
        if not os.path.exists(md_file):
            print(f"❌ 错误: 找不到文件 {md_file}")
            return
            
        try:
            create_word_document(md_file, output_file, template)
            print_success_info(output_file)
        except Exception as e:
            print(f"❌ 错误: {e}")
            
    elif len(sys.argv) == 2:
        # 用户只指定了输入文件
        md_file = sys.argv[1]
        if not os.path.isabs(md_file):
            alt = os.path.join(script_dir, md_file)
            if os.path.exists(alt):
                md_file = alt
        output_file = generate_output_filename(md_file)
        
        if not os.path.exists(md_file):
            print(f"❌ 错误: 找不到文件 {md_file}")
            return
            
        try:
            create_word_document(md_file, output_file, template_file)
            print_success_info(output_file)
        except Exception as e:
            print(f"❌ 错误: {e}")
            
    else:
        # 自动模式：处理当前目录下的所有.md文件
        md_files = find_md_files()
        
        if not md_files:
            print("❌ 当前目录下没有找到.md文件")
            print("\n💡 使用方法:")
            print("1. 将此脚本放在包含.md文件的文件夹中")
            print("2. 或者运行: python md_to_word_converter_complete.py 输入文件.md")
            print("3. 或者运行: python md_to_word_converter_complete.py 输入文件.md 输出文件.docx")
            print("4. 或者运行: python md_to_word_converter_complete.py 输入文件.md 输出文件.docx 模板文件.docx")
            print("\n📋 模板文件说明:")
            print("- 程序会自动查找同目录下的.docx文件作为模板")
            print("- 优先使用文件名包含'模板'或'template'的文件")
            print("- 使用模板时会保留模板的页面设置和样式，但清空内容")
            return
        
        print(f"🔍 找到 {len(md_files)} 个Markdown文件:")
        for i, md_file in enumerate(md_files, 1):
            print(f"  {i}. {md_file}")
        
        print("\n开始转换...")
        
        success_count = 0
        for md_file in md_files:
            output_file = generate_output_filename(md_file)
            try:
                create_word_document(md_file, output_file, template_file)
                success_count += 1
            except Exception as e:
                print(f"❌ 处理 {md_file} 时出错: {e}")
        
        print(f"\n✅ 转换完成！成功处理 {success_count}/{len(md_files)} 个文件")
        print_success_info()

def print_success_info(filename=None):
    """打印成功信息"""
    print("\n📋 自动应用的格式:")
    print("📄 页面大小: A4 (21cm × 29.7cm)")
    print("📝 字体: 仿宋_GB2312，黑色")
    print("📏 字号: 小四 (12pt)，一级标题小三 (15pt)")
    print("📐 行距: 1.5倍")
    print("📝 首行缩进: 2个字符（二级标题、三级标题、正文段落）")
    print("🎯 一级标题: 居中加粗，段前段后0.5行，不缩进")
    print("🔸 二级标题: 加粗，首行缩进，两端对齐")
    print("📄 正文段落: 两端对齐，首行缩进")
    print("📄 自动删除所有空行")
    print("📐 页边距: 上下2.54cm，左右3.18cm")
    print("💬 引号转换: 英文引号自动转为中文引号")
    print("📄 页码设置: 自动添加页脚页码（1/x格式，Times New Roman五号）")
    print("📊 表格支持: Markdown表格自动转换为Word表格，带边框格式")
    print("📈 图表支持: Mermaid图表本地渲染为高清图片插入Word文档（支持本地优先+在线备用）")
    print("✨ 格式支持: 支持**加粗**、*斜体*、<u>下划线</u>、~~删除线~~格式")
    print("🎯 表格内格式: 表格单元格内同样支持所有格式标记")
    print("\n🎯 完全无需手动调整！直接可用！")
    
    if filename:
        print(f"\n📁 输出文件: {filename}")

def debug_quotes_in_file(file_path):
    """简化的引号调试"""
    print("🔍 检查文件中的引号...")
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # 快速统计引号（区分ASCII与中文引号）
    ascii_double = content.count('"')
    chinese_open = content.count('“')
    chinese_close = content.count('”')

    print(f"📊 引号统计: ASCII双引号={ascii_double}, 中文开引号={chinese_open}, 中文闭引号={chinese_close}")

    # 只测试第一行包含引号的内容
    for i, line in enumerate(content.split('\n'), 1):
        if '"' in line:
            print(f"🎯 测试第{i}行: {line.strip()}")
            _ = convert_quotes_to_chinese(line.strip())
            break

    print("-" * 30)

if __name__ == "__main__":
    main() 

# 创建日期：250122 - 完整版本：支持表格转换和格式保留
